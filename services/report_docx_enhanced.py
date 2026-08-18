"""Enhanced Word exporter with professional diagrams and table-safe formatting.

This module is intentionally separate from the legacy exporter so the Report
Builder can be upgraded without destroying the existing fallback path.
"""
from __future__ import annotations
from io import BytesIO
import math
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from services.case03 import investor_benefits, investment_scenarios

FONT_NAME = "Times New Roman"
BODY_SIZE = 13
TABLE_SIZE = 9.5
CAPTION_SIZE = 10.5


def _field(p, instruction):
    run = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def _caption_text(text):
    text = str(text or "").strip()
    for prefix in ("Bảng ", "Hình "):
        if text.startswith(prefix):
            parts = text.split(". ", 1)
            if len(parts) == 2 and parts[0][len(prefix):].replace(".", "").isdigit():
                return parts[1]
    return text


def _add_caption(document, text, kind):
    p = document.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(); r.bold = True; r.font.name = FONT_NAME; r.font.size = Pt(CAPTION_SIZE)
    _field(p, f"SEQ {kind} \\* ARABIC")
    r = p.add_run(f". {_caption_text(text)}"); r.font.name = FONT_NAME; r.font.size = Pt(CAPTION_SIZE)


def _shade(cell, fill="DCE6F1"):
    tcpr = cell._tc.get_or_add_tcPr(); shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def _margins(cell, value=55):
    tcpr = cell._tc.get_or_add_tcPr(); mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar"); tcpr.append(mar)
    for side in ("top", "start", "bottom", "end"):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}"); mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def _repeat_header(row):
    trpr = row._tr.get_or_add_trPr(); node = OxmlElement("w:tblHeader"); node.set(qn("w:val"), "true"); trpr.append(node)


def _borders(table):
    tblpr = table._tbl.tblPr; borders = tblpr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders"); tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}"); borders.append(node)
        node.set(qn("w:val"), "single"); node.set(qn("w:sz"), "4"); node.set(qn("w:color"), "B7B7B7")


def _table_paragraph(p, center=False):
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fmt(value, key=""):
    if value is None: return ""
    if isinstance(value, bool): return "Có" if value else "Không"
    if isinstance(value, float):
        k = str(key).lower()
        if any(x in k for x in ("rate", "ratio", "margin", "ltv", "roi", "irr", "tỷ lệ", "lợi suất")):
            return f"{value * 100:,.2f}%"
        if abs(value) >= 1_000_000_000: return f"{value / 1_000_000_000:,.2f} tỷ đồng"
        if abs(value) < 10 and value != int(value): return f"{value:,.2f}"
        return f"{value:,.2f}"
    if isinstance(value, int): return f"{value:,}"
    return str(value)


def _records(data):
    if data is None: return []
    if hasattr(data, "to_dict"): data = data.to_dict("records")
    if not isinstance(data, list): data = [data]
    if data and not isinstance(data[0], dict): data = [{"Giá trị": x} for x in data]
    return data


def _set_table_widths(table, keys, rows, landscape=False):
    available = 24.5 if landscape else 16.2
    if len(keys) <= 2:
        ratios = [.28, .72]
    else:
        scores = []
        for key in keys:
            m = len(str(key))
            for row in rows[:40]:
                m = max(m, min(len(str(row.get(key, ""))), 42))
            scores.append(max(4, min(m, 32)))
        total = sum(scores); ratios = [x / total for x in scores]
    for row in table.rows:
        for i, ratio in enumerate(ratios):
            row.cells[i].width = Cm(max(1.25, available * ratio))


def add_table(document, data, title, caption, max_rows=300):
    rows = _records(data)
    if not rows: return None
    keys = []
    for row in rows:
        for key in row:
            if key not in keys: keys.append(key)
    wide = len(keys) >= 8
    if wide:
        section = document.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Cm(1.3); section.right_margin = Cm(1.3); section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
    h = document.add_heading(title, level=2); h.paragraph_format.first_line_indent = Cm(0)
    _add_caption(document, caption, "Table")
    table = document.add_table(rows=1, cols=len(keys)); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = True; _borders(table)
    header = table.rows[0]; _repeat_header(header)
    for i, key in enumerate(keys):
        cell = header.cells[i]; cell.text = str(key); _shade(cell); _margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            _table_paragraph(p, True)
            for r in p.runs: r.bold = True; r.font.name = FONT_NAME; r.font.size = Pt(8.5 if wide else TABLE_SIZE)
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for i, key in enumerate(keys):
            cell = cells[i]; cell.text = _fmt(row.get(key, ""), key); _margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                _table_paragraph(p, key in {"P", "I", "Điểm", "Thứ tự", "Bước"})
                for r in p.runs: r.font.name = FONT_NAME; r.font.size = Pt(8.5 if wide else TABLE_SIZE)
    _set_table_widths(table, keys, rows, wide)
    if wide:
        restore = document.add_section(WD_SECTION.NEW_PAGE)
        restore.orientation = WD_ORIENT.PORTRAIT
        restore.page_width, restore.page_height = restore.page_height, restore.page_width
        restore.top_margin = Cm(2.5); restore.bottom_margin = Cm(2.5); restore.left_margin = Cm(3); restore.right_margin = Cm(2)
    else:
        document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def _defaults(document):
    styles = document.styles; normal = styles["Normal"]
    normal.font.name = FONT_NAME; normal.font.size = Pt(BODY_SIZE); normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.paragraph_format.line_spacing = 1.35; normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.first_line_indent = Cm(1)
    for name, size in (("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13), ("Caption", CAPTION_SIZE)):
        s = styles[name]; s.font.name = FONT_NAME; s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME); s.font.size = Pt(size); s.paragraph_format.first_line_indent = Cm(0)
        if name.startswith("Heading"): s.font.bold = True; s.paragraph_format.space_before = Pt(12); s.paragraph_format.space_after = Pt(6)
    for section in document.sections:
        section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5); section.left_margin=Cm(3); section.right_margin=Cm(2); section.header_distance=Cm(1.2); section.footer_distance=Cm(1.2)
        header=section.header.paragraphs[0]; header.text="ĐỀ ÁN BLOCKCHAIN TRONG TÀI CHÍNH VÀ NGÂN HÀNG"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for r in header.runs:r.font.name=FONT_NAME; r.font.size=Pt(9)
        footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=footer.add_run("Trang "); rr.font.name=FONT_NAME; rr.font.size=Pt(9); _field(footer,"PAGE")
    for name in ("No Spacing", "Intense Quote"):
        if name in styles:
            styles[name].paragraph_format.first_line_indent=Cm(0); styles[name].paragraph_format.left_indent=Cm(0); styles[name].paragraph_format.right_indent=Cm(0)


def _list_field(document, title, instruction):
    document.add_heading(title, level=1); p=document.add_paragraph(); _table_paragraph(p); _field(p,instruction); note=document.add_paragraph("Word có thể cần cập nhật trường bằng Ctrl+A rồi F9."); _table_paragraph(note)
    for r in note.runs:r.italic=True; r.font.size=Pt(10)


def flow_figure(title, steps):
    steps=[str(x).strip() for x in (steps or []) if str(x).strip()] or ["Chưa có dữ liệu"]
    cols=2 if len(steps)>8 else 1; rows=math.ceil(len(steps)/cols)
    fig,ax=plt.subplots(figsize=(11,max(5.5,rows*1.25))); ax.set_xlim(0,1); ax.set_ylim(0,1); ax.axis("off"); ax.set_title(title,fontsize=15,fontweight="bold",pad=14)
    for i,step in enumerate(steps):
        col=i%cols; row=i//cols; x=.25 if cols==2 and col==0 else (.75 if cols==2 else .5); y=.90-row*(.76/max(1,rows-1)) if rows>1 else .52
        ax.text(x,y,f"{i+1}. {step}",ha="center",va="center",fontsize=9.2,wrap=True,bbox=dict(boxstyle="round,pad=.55",facecolor="#F4F6F8",edgecolor="#5B6573",linewidth=1.2))
        if i+1<len(steps):
            ni=i+1; nc=ni%cols; nr=ni//cols; nx=.25 if cols==2 and nc==0 else (.75 if cols==2 else .5); ny=.90-nr*(.76/max(1,rows-1)) if rows>1 else .52
            if cols==2 and col==0 and nc==1 and row==nr: ax.annotate("",xy=(nx-.08,ny),xytext=(x+.08,y),arrowprops=dict(arrowstyle="->",color="#667085",lw=1.1))
            else: ax.annotate("",xy=(nx,ny+.07),xytext=(x,y-.07),arrowprops=dict(arrowstyle="->",color="#667085",lw=1.1))
    fig.tight_layout(); out=BytesIO(); fig.savefig(out,format="png",dpi=190,bbox_inches="tight",facecolor="white"); plt.close(fig); out.seek(0); return out


def architecture_figure(case01):
    nodes=case01.get("architecture",{}).get("nodes",[]) or ["FutureBank"]
    fig,ax=plt.subplots(figsize=(11,6.4)); ax.set_xlim(0,1); ax.set_ylim(0,1); ax.axis("off"); ax.set_title("Kiến trúc Blockchain liên minh",fontsize=15,fontweight="bold",pad=12)
    ax.text(.5,.78,"SỔ CÁI LIÊN MINH\nBlockchain liên minh",ha="center",va="center",fontsize=13,fontweight="bold",bbox=dict(boxstyle="round,pad=.8",facecolor="#EAF0F6",edgecolor="#3F5B75",linewidth=1.4))
    positions=[(.15,.55),(.38,.55),(.62,.55),(.85,.55),(.25,.30),(.50,.30),(.75,.30),(.50,.10)]
    for i,node in enumerate(nodes[:len(positions)]):
        x,y=positions[i]; ax.plot([.5,x],[.72,y+.06],color="#98A2B3",linewidth=.9); ax.text(x,y,str(node),ha="center",va="center",fontsize=8.7,bbox=dict(boxstyle="round,pad=.45",facecolor="white",edgecolor="#667085",linewidth=1.0),wrap=True)
    ax.text(.12,.02,"Lớp dữ liệu ngoài chuỗi: KYC, hồ sơ chi tiết, chứng từ, báo cáo",ha="left",va="center",fontsize=8.5,bbox=dict(boxstyle="round,pad=.4",facecolor="#F7F7F7",edgecolor="#98A2B3")); ax.text(.88,.02,"Lớp trên chuỗi: trạng thái, bằng chứng, mã băm, giao dịch",ha="right",va="center",fontsize=8.5,bbox=dict(boxstyle="round,pad=.4",facecolor="#F7F7F7",edgecolor="#98A2B3"))
    fig.tight_layout(); out=BytesIO(); fig.savefig(out,format="png",dpi=190,bbox_inches="tight",facecolor="white"); plt.close(fig); out.seek(0); return out


def heatmap_figure(case01,case02,case03):
    points=[]
    for rows in (case01.get("risks",[]),case02.get("risks",[]),case03.get("risks",[])):
        for row in rows:
            try:p,impact=int(row.get("P",0)),int(row.get("I",0))
            except (TypeError,ValueError):continue
            if 1<=p<=5 and 1<=impact<=5:points.append((p,impact,str(row.get("Rủi ro","Rủi ro"))))
    fig,ax=plt.subplots(figsize=(9,6.8)); ax.set_xlim(.5,5.5); ax.set_ylim(.5,5.5); ax.set_xticks(range(1,6)); ax.set_yticks(range(1,6)); ax.set_xlabel("Tác động"); ax.set_ylabel("Xác suất"); ax.set_title("Integrated Risk Heat Map",fontsize=14,fontweight="bold"); ax.grid(True,linewidth=.5)
    for p,impact,name in points:ax.scatter(impact,p,s=75); ax.annotate(name,(impact,p),xytext=(4,4),textcoords="offset points",fontsize=7)
    fig.tight_layout(); out=BytesIO(); fig.savefig(out,format="png",dpi=190,bbox_inches="tight",facecolor="white"); plt.close(fig); out.seek(0); return out


def roadmap_figure():
    fig,ax=plt.subplots(figsize=(11,3.2)); ax.set_xlim(0,1); ax.set_ylim(0,1); ax.axis("off"); ax.set_title("Lộ trình triển khai",fontsize=14,fontweight="bold"); stages=[("Proof of Concept","Kiểm chứng kỹ thuật"),("Pilot","Giao dịch giới hạn"),("Triển khai chính thức","Mở rộng và vận hành")]; xs=[.18,.5,.82]
    for i,((name,desc),x) in enumerate(zip(stages,xs)):
        ax.text(x,.58,name,ha="center",va="center",fontsize=11,fontweight="bold",bbox=dict(boxstyle="round,pad=.6",facecolor="#F4F6F8",edgecolor="#5B6573")); ax.text(x,.27,desc,ha="center",va="center",fontsize=9)
        if i<2:ax.annotate("",xy=(xs[i+1]-.08,.58),xytext=(x+.08,.58),arrowprops=dict(arrowstyle="->",color="#667085",lw=1.2))
    fig.tight_layout(); out=BytesIO(); fig.savefig(out,format="png",dpi=190,bbox_inches="tight",facecolor="white"); plt.close(fig); out.seek(0); return out


def preview_figures(case01,case02,case03):
    return {
        "As-is Process": flow_figure("As-is Process",[x.get("Hành động","") for x in case01.get("as_is",[])]),
        "Kiến trúc Blockchain liên minh": architecture_figure(case01),
        "To-be Process": flow_figure("To-be Process",[x.get("Hành động","") for x in case02.get("to_be",[])]),
        "Vòng đời token": flow_figure("Vòng đời token",[x.get("Giai đoạn","") for x in case03.get("lifecycle",[])]),
        "Integrated Risk Heat Map": heatmap_figure(case01,case02,case03),
        "Lộ trình triển khai": roadmap_figure(),
    }


def _add_figure(document,image,caption):
    p=document.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.add_run().add_picture(image,width=Cm(16.2)); _add_caption(document,caption,"Figure")


def _term_sheet(case03):
    fields=[("Tên doanh nghiệp",case03.get("business_name")),("Tên dự án",case03.get("project_name")),("Tên token",case03.get("token_name")),("Mã token",case03.get("token_code")),("Công cụ huy động",case03.get("instrument")),("Tài sản cơ sở",case03.get("asset_base")),("Giá phát hành",case03.get("issue_price")),("Lợi suất giả định",case03.get("annual_return_rate")),("Nhà đầu tư mục tiêu",case03.get("target_investor")),("Hạn chế chuyển nhượng",case03.get("transfer_restrictions")),("Cơ chế lưu ký",case03.get("custody")),("Mua lại",case03.get("buyback")),("Thứ tự ưu tiên thanh toán",case03.get("payment_priority")),("Cấu trúc pháp lý",case03.get("legal_structure")),("Khoảng cách kỹ thuật và pháp lý",case03.get("legal_technical_gap"))]
    return [{"Trường":k,"Giá trị":_fmt(v,k)} for k,v in fields]


def _financial_summary(f):
    rows=[]
    for label,key in (("Tổng nhu cầu vốn","V"),("Khoản vay","LoanAmount"),("Vốn còn thiếu","ExternalCapital"),("Thời hạn","T"),("Lãi suất","r"),("Lãi vay năm đầu","Interest"),("Doanh thu năm đầu","Revenue1"),("EBITDA năm đầu","EBITDA1"),("Nghĩa vụ nợ năm đầu","DebtService1"),("DSCR","DSCR"),("Giá trị tài sản bảo đảm","CollateralValue"),("LTV","LTV"),("Dòng tiền còn lại","ResidualCash")):
        v=f.get(key,0)
        display=f"{float(v or 0)*100:,.2f}%" if key in {"r","LTV"} else (f"{float(v or 0):,.2f} lần" if key=="DSCR" else (f"{float(v or 0):,.0f} năm" if key=="T" else f"{float(v or 0):,.2f} tỷ đồng"))
        rows.append({"Chỉ tiêu":label,"Giá trị":display})
    return rows


def build_docx(profile,financials,case01,case02,case03,report,consistency_results,quality_checks=None):
    document=Document(); _defaults(document)
    p=document.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(75); p.paragraph_format.first_line_indent=Cm(0); r=p.add_run("ĐỀ ÁN XÂY DỰNG HỆ SINH THÁI BLOCKCHAIN\nTÍCH HỢP TÍN DỤNG, HUY ĐỘNG VỐN VÀ ĐẦU TƯ"); r.bold=True; r.font.size=Pt(20)
    p=document.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_before=Pt(40); r=p.add_run("BÁO CÁO TỔNG HỢP CASE 01 – CASE 02 – CASE 03"); r.bold=True; r.font.size=Pt(15)
    p=document.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_before=Pt(35); p.add_run(f"Sinh viên: {profile.get('student_id','')}\n").bold=True; p.add_run(f"Doanh nghiệp: {profile.get('business_name',profile.get('business_type',''))}\n"); p.add_run(f"Ngành: {profile.get('industry','')}\nFutureBank – Blockchain Finance Case Study"); document.add_page_break()
    _list_field(document,"Mục lục",'TOC \\o "1-3" \\h \\z \\u'); document.add_page_break(); _list_field(document,"Danh mục bảng",'TOC \\c "Table" \\h \\z \\u'); document.add_page_break(); _list_field(document,"Danh mục hình",'TOC \\c "Figure" \\h \\z \\u'); document.add_page_break()
    document.add_heading("Tóm tắt điều hành",level=1); document.add_paragraph(report.get("executive","").strip() or "Chưa có nội dung phân tích của sinh viên."); add_table(document,_financial_summary(financials),"Tóm tắt thông số tài chính","Tóm tắt thông số tài chính của case")
    document.add_page_break(); document.add_heading("Phần 1. Hồ sơ case cá nhân",level=1); document.add_paragraph(report.get("part1","").strip() or "Chưa có nội dung phân tích của sinh viên."); add_table(document,[{"Thông số":k,"Giá trị":profile.get(k,"")} for k in ("student_id","D1","D2","D3","D4","industry","business_type","banking_problem","funding_instrument","business_description")],"Hồ sơ doanh nghiệp","Bảng 1.1. Hồ sơ case cá nhân"); add_table(document,_financial_summary(financials),"Cấu trúc vốn","Bảng 1.2. Cấu trúc vốn của case")
    as_is=sorted(case01.get("as_is",[]),key=lambda x:int(x.get("Thứ tự",x.get("Bước",0)) or 0)); document.add_page_break(); document.add_heading("Phần 2. Phân tích vấn đề và quy trình hiện tại",level=1); document.add_paragraph(report.get("part2","").strip() or "Chưa có nội dung phân tích của sinh viên."); _add_figure(document,flow_figure("As-is Process",[x.get("Hành động","") for x in as_is]),"Hình 2.1. Quy trình As-is"); add_table(document,as_is,"As-is Process","Bảng 2.1. Quy trình As-is")
    stakeholders=[]; seen=set()
    for row in as_is:
        actor=row.get("Chủ thể","")
        if actor and actor not in seen:seen.add(actor); stakeholders.append({"Chủ thể":actor,"Vai trò":row.get("Trách nhiệm","")})
    add_table(document,stakeholders,"Các bên liên quan","Bảng 2.2. Các bên liên quan")
    document.add_page_break(); document.add_heading("Phần 3. Thiết kế kiến trúc blockchain",level=1); document.add_paragraph(report.get("part3","").strip() or "Chưa có nội dung phân tích của sinh viên."); _add_figure(document,architecture_figure(case01),"Hình 3.1. Kiến trúc Blockchain liên minh"); add_table(document,case01.get("assessment",[]),"Đánh giá CSDL và Blockchain/DLT","Bảng 3.1. Đánh giá CSDL và Blockchain/DLT"); add_table(document,case01.get("permissions",[]),"Ma trận quyền","Bảng 3.2. Ma trận quyền"); add_table(document,case01.get("data",[]),"Phân loại On-chain và Off-chain","Bảng 3.3. Phân loại dữ liệu On-chain và Off-chain"); add_table(document,[{"Quy tắc quản trị":k,"Nội dung":v} for k,v in (case01.get("governance",{}) or {}).items()],"Quản trị mạng","Bảng 3.4. Cơ chế quản trị mạng"); add_table(document,case01.get("risks",[]),"Risk Register Case 01","Bảng 3.5. Risk Register Case 01")
    document.add_page_break(); document.add_heading("Phần 4. Thiết kế sản phẩm tín dụng",level=1); document.add_paragraph(report.get("part4","").strip() or "Chưa có nội dung phân tích của sinh viên."); to_be=case02.get("to_be",[]); _add_figure(document,flow_figure("To-be Process",[x.get("Hành động","") for x in to_be]),"Hình 4.1. Quy trình To-be"); add_table(document,to_be,"To-be Process","Bảng 4.1. Quy trình To-be"); add_table(document,_financial_summary(financials),"Phân tích tín dụng","Bảng 4.2. Chỉ tiêu tín dụng"); add_table(document,case02.get("oracle",[]),"Oracle","Bảng 4.3. Danh mục Oracle"); add_table(document,[{"Kịch bản":x} for x in case02.get("scenarios",["Cơ sở","Thuận lợi","Bất lợi"])],"Ba kịch bản","Bảng 4.4. Ba kịch bản tín dụng"); add_table(document,case02.get("risks",[]),"Risk Register Case 02","Bảng 4.5. Risk Register Case 02")
    document.add_page_break(); document.add_heading("Phần 5. Thiết kế phương án huy động vốn",level=1); document.add_paragraph(report.get("part5","").strip() or "Chưa có nội dung phân tích của sinh viên."); add_table(document,_term_sheet(case03),"Term Sheet","Bảng 5.1. Term Sheet"); lifecycle=case03.get("lifecycle",[]); _add_figure(document,flow_figure("Vòng đời token",[x.get("Giai đoạn","") for x in lifecycle]),"Hình 5.1. Vòng đời token"); add_table(document,lifecycle,"Vòng đời token","Bảng 5.2. Vòng đời token")
    try:add_table(document,[{"Chỉ tiêu":k,"Giá trị":v} for k,v in investor_benefits(case03,financials).items()],"Lợi ích nhà đầu tư","Bảng 5.3. Lợi ích giả định của nhà đầu tư")
    except Exception:pass
    try:add_table(document,investment_scenarios(case03,financials),"Phân tích ba kịch bản","Bảng 5.4. Phân tích ba kịch bản huy động vốn")
    except Exception:add_table(document,[{"Kịch bản":x} for x in case03.get("scenarios",[])],"Phân tích ba kịch bản","Bảng 5.4. Phân tích ba kịch bản huy động vốn")
    document.add_heading("Hợp đồng thông minh",level=2); document.add_paragraph(case03.get("smart_contract","Chưa có pseudocode."),style="No Spacing"); add_table(document,case03.get("risks",[]),"Risk Register Case 03","Bảng 5.5. Risk Register Case 03")
    document.add_page_break(); document.add_heading("Phần 6. Quản trị rủi ro tích hợp",level=1); document.add_paragraph(report.get("part6","").strip() or "Chưa có nội dung phân tích của sinh viên."); _add_figure(document,heatmap_figure(case01,case02,case03),"Hình 6.1. Integrated Risk Heat Map"); combined=[]
    for source,rows in (("Case 01",case01.get("risks",[])),("Case 02",case02.get("risks",[])),("Case 03",case03.get("risks",[]))):
        for row in rows:item=dict(row); item["Nguồn"]=source; combined.append(item)
    add_table(document,combined,"Risk Register tích hợp","Bảng 6.1. Risk Register tích hợp")
    document.add_page_break(); document.add_heading("Phần 7. Đánh giá hiệu quả",level=1); document.add_paragraph(report.get("part7","").strip() or "Chưa có nội dung phân tích của sinh viên."); kpis=case03.get("kpis") or case02.get("kpis") or case01.get("kpis") or []
    if kpis:add_table(document,kpis,"KPI","Bảng 7.1. Bộ KPI đánh giá hiệu quả")
    else:add_table(document,[{"KPI":"Thời gian KYC","Cách đo":"Thời gian từ tiếp nhận đến xác minh","Nguồn":"Hệ thống KYC","Tần suất":"Hàng tháng"},{"KPI":"Thời gian phê duyệt","Cách đo":"Thời gian từ hồ sơ đủ đến phê duyệt","Nguồn":"Core tín dụng","Tần suất":"Hàng tháng"},{"KPI":"Thời gian giải ngân","Cách đo":"Thời gian từ phê duyệt đến giải ngân","Nguồn":"Core banking","Tần suất":"Hàng tháng"}],"KPI khung","Bảng 7.1. KPI khung cần hoàn thiện")
    document.add_page_break(); document.add_heading("Phần 8. Lộ trình triển khai",level=1); document.add_paragraph(report.get("part8","").strip() or "Chưa có nội dung phân tích của sinh viên."); _add_figure(document,roadmap_figure(),"Hình 8.1. Lộ trình triển khai"); add_table(document,[{"Giai đoạn":"Proof of Concept","Mục tiêu":"Kiểm chứng kỹ thuật","Phạm vi":"Dữ liệu và giao dịch mô phỏng","Điều kiện chuyển":"Kiểm chứng được kiến trúc và kiểm soát rủi ro"},{"Giai đoạn":"Pilot","Mục tiêu":"Kiểm chứng vận hành","Phạm vi":"Giao dịch giới hạn","Điều kiện chuyển":"Đạt KPI và yêu cầu pháp lý"},{"Giai đoạn":"Triển khai chính thức","Mục tiêu":"Mở rộng vận hành","Phạm vi":"Hệ sinh thái đầy đủ","Điều kiện chuyển":"Phê duyệt quản trị và sẵn sàng vận hành"}],"Lộ trình triển khai","Bảng 8.1. Lộ trình triển khai")
    document.add_page_break(); document.add_heading("Phần 9. Kết luận và khuyến nghị",level=1); document.add_paragraph(report.get("part9","").strip() or "Chưa có nội dung phân tích của sinh viên."); add_table(document,[{"Hạng mục":r.get("Hạng mục",""),"Trạng thái":r.get("Trạng thái",""),"Chi tiết":r.get("Chi tiết","")} for r in consistency_results or []],"Tổng hợp Consistency Checker","Bảng 9.1. Tổng hợp Consistency Checker")
    document.add_page_break(); document.add_heading("Phụ lục",level=1); document.add_heading("Phụ lục 1. Bảng tính",level=2); add_table(document,[{"Biến":k,"Giá trị":_fmt(v,k)} for k,v in financials.items()],"Các biến tính toán tài chính","Bảng P1.1. Các biến tính toán tài chính"); document.add_heading("Phụ lục 2. Pseudocode và Term Sheet",level=2); document.add_paragraph(case03.get("smart_contract","Chưa có pseudocode."),style="No Spacing"); add_table(document,_term_sheet(case03),"Term Sheet chi tiết","Bảng P2.1. Term Sheet chi tiết"); document.add_heading("Phụ lục 3. Risk Register",level=2); add_table(document,case01.get("risks",[]),"Risk Register Case 01","Bảng P3.1. Risk Register Case 01"); add_table(document,case02.get("risks",[]),"Risk Register Case 02","Bảng P3.2. Risk Register Case 02"); add_table(document,case03.get("risks",[]),"Risk Register Case 03","Bảng P3.3. Risk Register Case 03"); document.add_heading("Phụ lục 4. Change Log",level=2); add_table(document,case03.get("change_log",[]),"Change Log","Bảng P4.1. Change Log") if case03.get("change_log") else document.add_paragraph("Sinh viên bổ sung ngày, nội dung thay đổi, lý do và tác động tới Case sau theo Instruction File.",style="Intense Quote"); document.add_heading("Phụ lục 5. Consistency Checker",level=2); add_table(document,consistency_results or [],"Kết quả kiểm tra 22 câu","Bảng P5.1. Kết quả Consistency Checker"); document.add_heading("Phụ lục 6. Kiểm tra chất lượng báo cáo",level=2); add_table(document,quality_checks or [],"Quality Gate","Bảng P6.1. Kiểm tra chất lượng trước khi xuất")
    settings=document.settings.element; update=settings.find(qn("w:updateFields"))
    if update is None:update=OxmlElement("w:updateFields"); settings.append(update)
    update.set(qn("w:val"),"true"); output=BytesIO(); document.save(output); return output.getvalue()

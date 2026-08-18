"""Professional Word report generator.

The exporter builds native Word objects instead of converting Markdown. Tables,
captions, headings, page numbers, figures, TOC and lists are real Word fields.
"""
from __future__ import annotations

from io import BytesIO
import math

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from services.case03 import investor_benefits, investment_scenarios
from services.report_builder import REPORT_SECTIONS


FONT_NAME = "Times New Roman"
BODY_SIZE = 13
TABLE_SIZE = 10
CAPTION_SIZE = 10.5


def _set_cell_shading(cell, fill="D9E2F3"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "B7B7B7")


def _set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def _add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def _add_caption(document, text, kind):
    p = document.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = Pt(CAPTION_SIZE)
    _add_field(p, f'SEQ {kind} \\* ARABIC')
    p.add_run(f". {text}")
    _set_keep_with_next(p)
    return p


def _fmt_value(value, key=""):
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Có" if value else "Không"
    if isinstance(value, float):
        k = str(key).lower()
        if any(x in k for x in ("rate", "ratio", "margin", "ltv", "roi", "irr", "tỷ lệ", "lợi suất")):
            return f"{value * 100:,.2f}%"
        if abs(value) >= 1_000_000_000:
            return f"{value / 1_000_000_000:,.2f} tỷ đồng"
        if abs(value) >= 1_000_000:
            return f"{value:,.0f}"
        if abs(value) < 10 and value != int(value):
            return f"{value:,.2f}"
        return f"{value:,.2f}"
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def _records(records):
    if records is None:
        return []
    if hasattr(records, "to_dict"):
        records = records.to_dict("records")
    if not isinstance(records, list):
        records = [records]
    if records and not isinstance(records[0], dict):
        records = [{"Giá trị": x} for x in records]
    return records


def _add_table(document, records, title, caption=None, max_rows=300):
    rows = _records(records)
    if not rows:
        return None
    keys = []
    for row in rows:
        for key in row.keys():
            if key not in keys:
                keys.append(key)
    if not keys:
        return None
    # Keep every field. Very wide tables use compact type instead of dropping data.
    rows = rows[:max_rows]
    p = document.add_heading(title, level=2)
    _set_keep_with_next(p)
    table = document.add_table(rows=1, cols=len(keys))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table)
    header = table.rows[0]
    _repeat_table_header(header)
    for i, key in enumerate(keys):
        cell = header.cells[i]
        cell.text = str(key)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell)
        _set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = FONT_NAME
                r.font.size = Pt(TABLE_SIZE)
    for row in rows:
        cells = table.add_row().cells
        for i, key in enumerate(keys):
            cell = cells[i]
            cell.text = _fmt_value(row.get(key, ""), key)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if key in {"P", "I", "Điểm", "Thứ tự", "Bước"}:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = FONT_NAME
                    r.font.size = Pt(TABLE_SIZE)
    _add_caption(document, caption or title, "Table")
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def _set_document_defaults(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1)

    for name, size in (("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13), ("Caption", CAPTION_SIZE)):
        style = styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        if name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)

    for section in document.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)


def _add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "ĐỀ ÁN BLOCKCHAIN TRONG TÀI CHÍNH VÀ NGÂN HÀNG"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.size = Pt(9)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Trang ")
    r.font.name = FONT_NAME
    r.font.size = Pt(9)
    _add_field(p, "PAGE")


def _add_toc(document, title, instruction):
    document.add_heading(title, level=1)
    p = document.add_paragraph()
    _add_field(p, instruction)
    document.add_paragraph("Nhấn Ctrl+A rồi F9 trong Word để cập nhật mục lục, danh mục bảng và danh mục hình nếu Word chưa tự cập nhật.").italic = True


def _flow_figure(title, steps, filename=None):
    fig, ax = plt.subplots(figsize=(11, 5.2))
    ax.axis("off")
    if not steps:
        steps = ["Chưa có dữ liệu"]
    n = len(steps)
    for i, step in enumerate(steps):
        x = (i + 0.5) / n
        ax.text(x, 0.5, str(step), ha="center", va="center", wrap=True, fontsize=9,
                bbox=dict(boxstyle="round,pad=0.5", facecolor="white", edgecolor="black"))
        if i < n - 1:
            ax.annotate("", xy=((i + 1) / n - 0.03, 0.5), xytext=(i / n + 0.03, 0.5),
                        arrowprops=dict(arrowstyle="->", lw=1.2))
    ax.set_title(title, fontsize=13, fontweight="bold")
    fig.tight_layout()
    out = BytesIO(); fig.savefig(out, format="png", dpi=180, bbox_inches="tight"); plt.close(fig); out.seek(0)
    return out


def _architecture_figure(case01):
    arch = case01.get("architecture", {})
    nodes = arch.get("nodes", []) or ["FutureBank"]
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.axis("off")
    ax.text(0.5, 0.78, "Blockchain liên minh", ha="center", va="center", fontsize=15, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.7", facecolor="white", edgecolor="black"))
    positions = []
    for i, node in enumerate(nodes):
        angle = 2 * math.pi * i / max(1, len(nodes))
        x = 0.5 + 0.36 * math.cos(angle)
        y = 0.42 + 0.28 * math.sin(angle)
        positions.append((x, y))
        ax.text(x, y, str(node), ha="center", va="center", fontsize=8, wrap=True,
                bbox=dict(boxstyle="round,pad=0.45", facecolor="white", edgecolor="black"))
        ax.annotate("", xy=(0.5 + 0.10 * (x - 0.5), 0.78 + 0.08 * (y - 0.78)), xytext=(x, y),
                    arrowprops=dict(arrowstyle="-", lw=0.8))
    ax.text(0.18, 0.08, "On-chain: trạng thái, bằng chứng, mã băm, giao dịch", fontsize=9,
            bbox=dict(boxstyle="round,pad=0.4", facecolor="white", edgecolor="black"))
    ax.text(0.62, 0.08, "Off-chain: KYC, hồ sơ chi tiết, báo cáo, chứng từ", fontsize=9,
            bbox=dict(boxstyle="round,pad=0.4", facecolor="white", edgecolor="black"))
    fig.tight_layout()
    out = BytesIO(); fig.savefig(out, format="png", dpi=180, bbox_inches="tight"); plt.close(fig); out.seek(0)
    return out


def _risk_heatmap(case01, case02, case03):
    risks = []
    for source, rows in (("Case 01", case01.get("risks", [])), ("Case 02", case02.get("risks", [])), ("Case 03", case03.get("risks", []))):
        for row in rows:
            try:
                p, impact = int(row.get("P", 0)), int(row.get("I", 0))
            except (TypeError, ValueError):
                continue
            if 1 <= p <= 5 and 1 <= impact <= 5:
                risks.append((p, impact, str(row.get("Rủi ro", "Rủi ro")), source))
    fig, ax = plt.subplots(figsize=(9, 7))
    ax.set_xlim(0.5, 5.5); ax.set_ylim(0.5, 5.5)
    ax.set_xticks(range(1, 6)); ax.set_yticks(range(1, 6))
    ax.set_xlabel("Tác động")
    ax.set_ylabel("Xác suất")
    ax.set_title("Integrated Risk Heat Map", fontsize=14, fontweight="bold")
    for p, impact, name, source in risks:
        ax.scatter(impact, p, s=100, marker="o")
        ax.annotate(name, (impact, p), xytext=(5, 4), textcoords="offset points", fontsize=7)
    ax.grid(True, linewidth=0.5)
    fig.tight_layout()
    out = BytesIO(); fig.savefig(out, format="png", dpi=180, bbox_inches="tight"); plt.close(fig); out.seek(0)
    return out


def _roadmap_figure():
    fig, ax = plt.subplots(figsize=(11, 3.5))
    ax.axis("off")
    stages = [("Proof of Concept", "Kiểm chứng kỹ thuật"), ("Pilot", "Giao dịch giới hạn"), ("Triển khai chính thức", "Mở rộng và vận hành")]
    xs = [0.17, 0.5, 0.83]
    for i, ((name, desc), x) in enumerate(zip(stages, xs)):
        ax.text(x, 0.55, name, ha="center", va="center", fontsize=12, fontweight="bold",
                bbox=dict(boxstyle="round,pad=0.6", facecolor="white", edgecolor="black"))
        ax.text(x, 0.25, desc, ha="center", va="center", fontsize=9)
        if i < 2:
            ax.annotate("", xy=(xs[i + 1] - 0.08, 0.55), xytext=(x + 0.08, 0.55),
                        arrowprops=dict(arrowstyle="->", lw=1.4))
    fig.tight_layout()
    out = BytesIO(); fig.savefig(out, format="png", dpi=180, bbox_inches="tight"); plt.close(fig); out.seek(0)
    return out


def _add_figure(document, image, caption):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image, width=Cm(16.2))
    _add_caption(document, caption, "Figure")


def _term_sheet(case03):
    fields = [
        ("Tên doanh nghiệp", case03.get("business_name")), ("Tên dự án", case03.get("project_name")),
        ("Tên token", case03.get("token_name")), ("Mã token", case03.get("token_code")),
        ("Công cụ huy động", case03.get("instrument")), ("Tài sản cơ sở", case03.get("asset_base")),
        ("Giá phát hành", case03.get("issue_price")), ("Lợi suất giả định", case03.get("annual_return_rate")),
        ("Nhà đầu tư mục tiêu", case03.get("target_investor")), ("Hạn chế chuyển nhượng", case03.get("transfer_restrictions")),
        ("Cơ chế lưu ký", case03.get("custody")), ("Mua lại", case03.get("buyback")),
        ("Thứ tự ưu tiên thanh toán", case03.get("payment_priority")), ("Cấu trúc pháp lý", case03.get("legal_structure")),
        ("Khoảng cách kỹ thuật và pháp lý", case03.get("legal_technical_gap")),
    ]
    return [{"Trường": k, "Giá trị": _fmt_value(v, k)} for k, v in fields]


def _financial_summary(financials):
    keys = [
        ("Tổng nhu cầu vốn", "V", "tỷ đồng"), ("Khoản vay", "LoanAmount", "tỷ đồng"),
        ("Vốn còn thiếu", "ExternalCapital", "tỷ đồng"), ("Thời hạn", "T", "năm"),
        ("Lãi suất", "r", "%"), ("Lãi vay năm đầu", "Interest", "tỷ đồng"),
        ("Doanh thu năm đầu", "Revenue1", "tỷ đồng"), ("EBITDA năm đầu", "EBITDA1", "tỷ đồng"),
        ("Nghĩa vụ nợ năm đầu", "DebtService1", "tỷ đồng"), ("DSCR", "DSCR", "lần"),
        ("Giá trị tài sản bảo đảm", "CollateralValue", "tỷ đồng"), ("LTV", "LTV", "%"),
        ("Dòng tiền còn lại", "ResidualCash", "tỷ đồng"),
    ]
    rows = []
    for label, key, unit in keys:
        value = financials.get(key)
        if key in {"r", "LTV"}:
            display = f"{float(value or 0) * 100:,.2f}%"
        elif key == "DSCR":
            display = f"{float(value or 0):,.2f} lần"
        else:
            display = f"{float(value or 0):,.2f} {unit}"
        rows.append({"Chỉ tiêu": label, "Giá trị": display})
    return rows


def build_docx(profile, financials, case01, case02, case03, report, consistency_results, quality_checks=None):
    document = Document()
    _set_document_defaults(document)
    for section in document.sections:
        _add_header_footer(section)

    # Cover
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("ĐỀ ÁN XÂY DỰNG HỆ SINH THÁI BLOCKCHAIN\nTÍCH HỢP TÍN DỤNG, HUY ĐỘNG VỐN VÀ ĐẦU TƯ")
    r.bold = True; r.font.name = FONT_NAME; r.font.size = Pt(20)
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    r = p.add_run("BÁO CÁO TỔNG HỢP CASE 01 – CASE 02 – CASE 03")
    r.bold = True; r.font.size = Pt(15)
    document.add_paragraph("")
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Sinh viên: {profile.get('student_id', '')}\n").bold = True
    p.add_run(f"Doanh nghiệp: {profile.get('business_name', profile.get('business_type', ''))}\n")
    p.add_run(f"Ngành: {profile.get('industry', '')}\n")
    p.add_run("FutureBank – Blockchain Finance Case Study")
    document.add_page_break()

    _add_toc(document, "Mục lục", 'TOC \\o "1-3" \\h \\z \\u')
    document.add_page_break()
    _add_toc(document, "Danh mục bảng", 'TOC \\c "Table" \\h \\z \\u')
    document.add_page_break()
    _add_toc(document, "Danh mục hình", 'TOC \\c "Figure" \\h \\z \\u')
    document.add_page_break()

    # Executive summary
    document.add_heading("Tóm tắt điều hành", level=1)
    executive = report.get("executive", "").strip()
    document.add_paragraph(executive or "Chưa có nội dung phân tích của sinh viên.")
    _add_table(document, _financial_summary(financials), "Tóm tắt các thông số tài chính", "Tóm tắt thông số tài chính của case")

    # Part 1
    document.add_page_break(); document.add_heading("Phần 1. Hồ sơ case cá nhân", level=1)
    document.add_paragraph(report.get("part1", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    profile_rows = [
        {"Thông số": "MSSV", "Giá trị": profile.get("student_id", "")},
        {"Thông số": "D1", "Giá trị": profile.get("D1", "")}, {"Thông số": "D2", "Giá trị": profile.get("D2", "")},
        {"Thông số": "D3", "Giá trị": profile.get("D3", "")}, {"Thông số": "D4", "Giá trị": profile.get("D4", "")},
        {"Thông số": "Ngành hoạt động", "Giá trị": profile.get("industry", "")},
        {"Thông số": "Loại hình doanh nghiệp", "Giá trị": profile.get("business_type", "")},
        {"Thông số": "Vấn đề ngân hàng trọng tâm", "Giá trị": profile.get("banking_problem", "")},
        {"Thông số": "Công cụ huy động vốn Case 03", "Giá trị": profile.get("funding_instrument", "")},
        {"Thông số": "Mô tả hoạt động", "Giá trị": profile.get("business_description", "")},
    ]
    _add_table(document, profile_rows, "Hồ sơ doanh nghiệp", "Bảng 1.1. Hồ sơ case cá nhân")
    _add_table(document, _financial_summary(financials), "Cấu trúc vốn", "Bảng 1.2. Cấu trúc vốn của case")

    # Part 2
    document.add_page_break(); document.add_heading("Phần 2. Phân tích vấn đề và quy trình hiện tại", level=1)
    document.add_paragraph(report.get("part2", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    as_is = sorted(case01.get("as_is", []), key=lambda x: int(x.get("Thứ tự", x.get("Bước", 0)) or 0))
    _add_figure(document, _flow_figure("As-is Process", [x.get("Hành động", "") for x in as_is]), "Hình 2.1. Quy trình As-is")
    _add_table(document, as_is, "As-is Process", "Bảng 2.1. Quy trình As-is")
    stakeholders = []
    seen = set()
    for row in as_is:
        actor = row.get("Chủ thể", "")
        if actor and actor not in seen:
            seen.add(actor); stakeholders.append({"Chủ thể": actor, "Vai trò": row.get("Trách nhiệm", "")})
    _add_table(document, stakeholders, "Các bên liên quan", "Bảng 2.2. Các bên liên quan")

    # Part 3
    document.add_page_break(); document.add_heading("Phần 3. Thiết kế kiến trúc blockchain", level=1)
    document.add_paragraph(report.get("part3", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    arch = case01.get("architecture", {})
    _add_figure(document, _architecture_figure(case01), "Hình 3.1. Kiến trúc Blockchain liên minh")
    _add_table(document, case01.get("assessment", []), "Đánh giá CSDL và Blockchain/DLT", "Bảng 3.1. Đánh giá CSDL và Blockchain/DLT")
    _add_table(document, case01.get("permissions", []), "Ma trận quyền", "Bảng 3.2. Ma trận quyền")
    _add_table(document, case01.get("data", []), "Phân loại On-chain và Off-chain", "Bảng 3.3. Phân loại dữ liệu On-chain và Off-chain")
    governance_rows = [{"Quy tắc quản trị": k, "Nội dung": v} for k, v in (case01.get("governance", {}) or {}).items()]
    _add_table(document, governance_rows, "Quản trị mạng", "Bảng 3.4. Cơ chế quản trị mạng")
    _add_table(document, case01.get("risks", []), "Risk Register Case 01", "Bảng 3.5. Risk Register Case 01")
    _add_table(document, [{"Thông số": "Quyết định", "Giá trị": arch.get("decision")}, {"Thông số": "Loại Blockchain", "Giá trị": arch.get("blockchain_type")}, {"Thông số": "Đồng thuận", "Giá trị": arch.get("consensus")}, {"Thông số": "Số validator", "Giá trị": arch.get("validator_count")}, {"Thông số": "Điều kiện hoàn tất", "Giá trị": arch.get("completion")}], "Thông số kiến trúc", "Bảng 3.6. Thông số kiến trúc")

    # Part 4
    document.add_page_break(); document.add_heading("Phần 4. Thiết kế sản phẩm tín dụng", level=1)
    document.add_paragraph(report.get("part4", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    to_be = case02.get("to_be", [])
    _add_figure(document, _flow_figure("To-be Process", [x.get("Hành động", "") for x in to_be]), "Hình 4.1. Quy trình To-be")
    _add_table(document, to_be, "To-be Process", "Bảng 4.1. Quy trình To-be")
    _add_table(document, _financial_summary(financials), "Phân tích tín dụng", "Bảng 4.2. Chỉ tiêu tín dụng")
    _add_table(document, case02.get("oracle", []), "Oracle", "Bảng 4.3. Danh mục Oracle")
    _add_table(document, [{"Kịch bản": "Cơ sở"}, {"Kịch bản": "Thuận lợi"}, {"Kịch bản": "Bất lợi"}], "Ba kịch bản", "Bảng 4.4. Ba kịch bản tín dụng")
    _add_table(document, case02.get("risks", []), "Risk Register Case 02", "Bảng 4.5. Risk Register Case 02")
    _add_table(document, case02.get("smart_contract_events", []), "Điều kiện hợp đồng thông minh", "Bảng 4.6. Các điều kiện kích hoạt hợp đồng thông minh")

    # Part 5
    document.add_page_break(); document.add_heading("Phần 5. Thiết kế phương án huy động vốn", level=1)
    document.add_paragraph(report.get("part5", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    _add_table(document, _term_sheet(case03), "Term Sheet", "Bảng 5.1. Term Sheet")
    lifecycle = case03.get("lifecycle", [])
    _add_figure(document, _flow_figure("Vòng đời token", [x.get("Giai đoạn", "") for x in lifecycle]), "Hình 5.1. Vòng đời token")
    _add_table(document, lifecycle, "Vòng đời token", "Bảng 5.2. Vòng đời token")
    try:
        benefit = investor_benefits(case03, financials)
        benefit_rows = [{"Chỉ tiêu": k, "Giá trị": v} for k, v in benefit.items()]
        _add_table(document, benefit_rows, "Lợi ích nhà đầu tư", "Bảng 5.3. Lợi ích giả định của nhà đầu tư")
    except Exception:
        pass
    try:
        scenarios = investment_scenarios(case03, financials)
        _add_table(document, scenarios, "Phân tích ba kịch bản", "Bảng 5.4. Phân tích ba kịch bản huy động vốn")
    except Exception:
        _add_table(document, [{"Kịch bản": x} for x in case03.get("scenarios", [])], "Phân tích ba kịch bản", "Bảng 5.4. Phân tích ba kịch bản huy động vốn")
    document.add_heading("Hợp đồng thông minh", level=2)
    document.add_paragraph(case03.get("smart_contract", "Chưa có pseudocode."), style="No Spacing")
    _add_table(document, case03.get("risks", []), "Risk Register Case 03", "Bảng 5.5. Risk Register Case 03")

    # Part 6
    document.add_page_break(); document.add_heading("Phần 6. Quản trị rủi ro tích hợp", level=1)
    document.add_paragraph(report.get("part6", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    _add_figure(document, _risk_heatmap(case01, case02, case03), "Hình 6.1. Integrated Risk Heat Map")
    combined = []
    for source, rows in (("Case 01", case01.get("risks", [])), ("Case 02", case02.get("risks", [])), ("Case 03", case03.get("risks", []))):
        for row in rows:
            item = dict(row); item["Nguồn"] = source; combined.append(item)
    _add_table(document, combined, "Risk Register tích hợp", "Bảng 6.1. Risk Register tích hợp")

    # Part 7
    document.add_page_break(); document.add_heading("Phần 7. Đánh giá hiệu quả", level=1)
    document.add_paragraph(report.get("part7", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    kpis = case03.get("kpis") or case02.get("kpis") or case01.get("kpis") or []
    if kpis:
        _add_table(document, kpis, "KPI", "Bảng 7.1. Bộ KPI đánh giá hiệu quả")
    else:
        _add_table(document, [{"KPI": "Thời gian KYC", "Cách đo": "Thời gian từ tiếp nhận đến xác minh", "Nguồn": "Hệ thống KYC", "Tần suất": "Hàng tháng"}, {"KPI": "Thời gian phê duyệt", "Cách đo": "Thời gian từ hồ sơ đủ đến phê duyệt", "Nguồn": "Core tín dụng", "Tần suất": "Hàng tháng"}, {"KPI": "Thời gian giải ngân", "Cách đo": "Thời gian từ phê duyệt đến giải ngân", "Nguồn": "Core banking", "Tần suất": "Hàng tháng"}], "KPI khung", "Bảng 7.1. KPI khung cần hoàn thiện")
        document.add_paragraph("Bảng KPI khung được đưa vào để tránh báo cáo bị thiếu cấu trúc. Sinh viên cần hoàn thiện tối thiểu 10 KPI theo Instruction File trước khi nộp.", style="Intense Quote")

    # Part 8
    document.add_page_break(); document.add_heading("Phần 8. Lộ trình triển khai", level=1)
    document.add_paragraph(report.get("part8", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    _add_figure(document, _roadmap_figure(), "Hình 8.1. Lộ trình triển khai")
    roadmap = [{"Giai đoạn": "Proof of Concept", "Mục tiêu": "Kiểm chứng kỹ thuật", "Phạm vi": "Dữ liệu và giao dịch mô phỏng", "Điều kiện chuyển": "Kiểm chứng được kiến trúc và kiểm soát rủi ro"}, {"Giai đoạn": "Pilot", "Mục tiêu": "Kiểm chứng vận hành", "Phạm vi": "Giao dịch giới hạn", "Điều kiện chuyển": "Đạt KPI và yêu cầu pháp lý"}, {"Giai đoạn": "Triển khai chính thức", "Mục tiêu": "Mở rộng vận hành", "Phạm vi": "Hệ sinh thái đầy đủ", "Điều kiện chuyển": "Phê duyệt quản trị và sẵn sàng vận hành"}]
    _add_table(document, roadmap, "Lộ trình triển khai", "Bảng 8.1. Lộ trình triển khai")

    # Part 9
    document.add_page_break(); document.add_heading("Phần 9. Kết luận và khuyến nghị", level=1)
    document.add_paragraph(report.get("part9", "").strip() or "Chưa có nội dung phân tích của sinh viên.")
    _add_table(document, [{"Hạng mục": r.get("Hạng mục", ""), "Trạng thái": r.get("Trạng thái", ""), "Chi tiết": r.get("Chi tiết", "")} for r in consistency_results or []], "Tổng hợp Consistency Checker", "Bảng 9.1. Tổng hợp Consistency Checker")

    # Appendices
    document.add_page_break(); document.add_heading("Phụ lục", level=1)
    document.add_heading("Phụ lục 1. Bảng tính", level=2)
    _add_table(document, [{"Biến": k, "Giá trị": _fmt_value(v, k)} for k, v in financials.items()], "Các biến tính toán tài chính", "Bảng P1.1. Các biến tính toán tài chính")
    document.add_heading("Phụ lục 2. Pseudocode và Term Sheet", level=2)
    document.add_paragraph(case03.get("smart_contract", "Chưa có pseudocode."), style="No Spacing")
    _add_table(document, _term_sheet(case03), "Term Sheet chi tiết", "Bảng P2.1. Term Sheet chi tiết")
    document.add_heading("Phụ lục 3. Risk Register", level=2)
    _add_table(document, case01.get("risks", []), "Risk Register Case 01", "Bảng P3.1. Risk Register Case 01")
    _add_table(document, case02.get("risks", []), "Risk Register Case 02", "Bảng P3.2. Risk Register Case 02")
    _add_table(document, case03.get("risks", []), "Risk Register Case 03", "Bảng P3.3. Risk Register Case 03")
    document.add_heading("Phụ lục 4. Change Log", level=2)
    _add_table(document, case03.get("change_log", []), "Change Log", "Bảng P4.1. Change Log") if case03.get("change_log") else document.add_paragraph("Sinh viên bổ sung ngày, nội dung thay đổi, lý do và tác động tới Case sau theo Instruction File.", style="Intense Quote")
    document.add_heading("Phụ lục 5. Consistency Checker", level=2)
    _add_table(document, consistency_results or [], "Kết quả kiểm tra 22 câu", "Bảng P5.1. Kết quả Consistency Checker")
    document.add_heading("Phụ lục 6. Kiểm tra chất lượng báo cáo", level=2)
    if quality_checks:
        _add_table(document, quality_checks, "Quality Gate", "Bảng P6.1. Kiểm tra chất lượng trước khi xuất")

    # Final field update settings
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output = BytesIO()
    document.save(output)
    return output.getvalue()
